#!/usr/bin/env python3
"""Convert SESSION_SUMMARY.md to a formatted .docx document.

Parses markdown headings, bullet lists, bold text, and horizontal rules,
then writes a styled Word document suitable for external communication.

Usage:
    python scripts/session_to_docx.py
    python scripts/session_to_docx.py --source SESSION_SUMMARY.md --output out.docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import List, Tuple


def _add_inline_bold(paragraph, text: str) -> None:
    """Write text into *paragraph*, converting **bold** markers to bold runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def _parse_lines(lines: List[str]) -> List[Tuple[str, str]]:
    """Convert raw markdown lines into (kind, content) tokens.

    Kinds: heading1, heading2, heading3, bullet, hr, code, blank, paragraph.
    """
    tokens: List[Tuple[str, str]] = []
    in_code = False
    for line in lines:
        stripped = line.rstrip("\n")
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            tokens.append(("code", stripped))
            continue
        if re.match(r"^#{1}\s", stripped):
            tokens.append(("heading1", stripped.lstrip("# ").strip()))
        elif re.match(r"^#{2}\s", stripped):
            tokens.append(("heading2", stripped.lstrip("# ").strip()))
        elif re.match(r"^#{3}\s", stripped):
            tokens.append(("heading3", stripped.lstrip("# ").strip()))
        elif re.match(r"^[-*]\s", stripped):
            tokens.append(("bullet", stripped[2:].strip()))
        elif re.match(r"^---+$", stripped):
            tokens.append(("hr", ""))
        elif stripped == "":
            tokens.append(("blank", ""))
        else:
            tokens.append(("paragraph", stripped))
    return tokens


def convert(source: Path, output: Path) -> None:
    """Read *source* markdown and write a formatted .docx to *output*."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print(
            "python-docx is required. Install with:\n"
            "  pip install python-docx\n"
            "  pip install -e .[sync]",
            file=sys.stderr,
        )
        sys.exit(1)

    if not source.exists():
        print(f"Source file not found: {source}", file=sys.stderr)
        sys.exit(1)

    lines = source.read_text(encoding="utf-8").splitlines(keepends=True)
    tokens = _parse_lines(lines)

    doc = Document()

    # Narrow default margins for a cleaner look
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(80)
        section.right_margin = Pt(80)

    for kind, content in tokens:
        if kind == "heading1":
            h = doc.add_heading(level=1)
            h.clear()
            run = h.add_run(content)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif kind == "heading2":
            h = doc.add_heading(level=2)
            h.clear()
            run = h.add_run(content)
            run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
        elif kind == "heading3":
            h = doc.add_heading(level=3)
            h.clear()
            run = h.add_run(content)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(p, content)
        elif kind == "hr":
            # Horizontal rule rendered as a thin paragraph border
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            _add_inline_bold(p, content)
        # blank lines are intentionally skipped (Word handles spacing via styles)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Saved: {output}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert SESSION_SUMMARY.md to a formatted .docx"
    )
    parser.add_argument(
        "--source",
        default="SESSION_SUMMARY.md",
        help="Path to input markdown file (default: SESSION_SUMMARY.md)",
    )
    parser.add_argument(
        "--output",
        default="outputs/SESSION_SUMMARY.docx",
        help="Path for output .docx (default: outputs/SESSION_SUMMARY.docx)",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    convert(Path(args.source), Path(args.output))
