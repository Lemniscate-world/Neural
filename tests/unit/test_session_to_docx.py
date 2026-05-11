"""
Unit tests for scripts/session_to_docx.py (NDBG-5).

Verifies markdown parsing, docx generation, and cross-platform path handling.
"""

import sys
from pathlib import Path

import pytest

# Make the infrastructure/scripts/ directory importable
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "infrastructure" / "scripts"))

from session_to_docx import _parse_lines, _add_inline_bold, convert


class TestParseLines:
    """Tokeniser covers all markdown constructs present in SESSION_SUMMARY.md."""

    def test_heading1(self):
        tokens = _parse_lines(["# My Title\n"])
        assert tokens == [("heading1", "My Title")]

    def test_heading2(self):
        tokens = _parse_lines(["## Section\n"])
        assert tokens == [("heading2", "Section")]

    def test_heading3(self):
        tokens = _parse_lines(["### Sub\n"])
        assert tokens == [("heading3", "Sub")]

    def test_bullet_dash(self):
        tokens = _parse_lines(["- item one\n"])
        assert tokens == [("bullet", "item one")]

    def test_bullet_star(self):
        tokens = _parse_lines(["* item two\n"])
        assert tokens == [("bullet", "item two")]

    def test_horizontal_rule(self):
        tokens = _parse_lines(["---\n"])
        assert tokens == [("hr", "")]

    def test_blank_line(self):
        tokens = _parse_lines(["\n"])
        assert tokens == [("blank", "")]

    def test_plain_paragraph(self):
        tokens = _parse_lines(["Hello world\n"])
        assert tokens == [("paragraph", "Hello world")]

    def test_code_block_skipped(self):
        lines = ["```python\n", "x = 1\n", "```\n"]
        tokens = _parse_lines(lines)
        assert ("code", "x = 1") in tokens
        # fence markers themselves are consumed
        kinds = [k for k, _ in tokens]
        assert "heading1" not in kinds

    def test_mixed_content(self):
        lines = [
            "# Title\n",
            "\n",
            "## English\n",
            "- bullet\n",
            "---\n",
        ]
        tokens = _parse_lines(lines)
        assert tokens[0] == ("heading1", "Title")
        assert tokens[1] == ("blank", "")
        assert tokens[2] == ("heading2", "English")
        assert tokens[3] == ("bullet", "bullet")
        assert tokens[4] == ("hr", "")


class TestInlineBold:
    """_add_inline_bold splits text on **...** markers and sets bold correctly."""

    def test_plain_text(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        _add_inline_bold(p, "plain text")
        runs = p.runs
        assert len(runs) == 1
        assert runs[0].text == "plain text"
        assert not runs[0].bold

    def test_bold_only(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        _add_inline_bold(p, "**bold**")
        runs = p.runs
        assert any(r.bold and r.text == "bold" for r in runs)

    def test_mixed_bold(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        _add_inline_bold(p, "before **bold** after")
        texts = [r.text for r in p.runs]
        assert "before " in texts
        assert "bold" in texts
        assert " after" in texts
        bold_runs = [r for r in p.runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"


class TestConvert:
    """End-to-end: convert() writes a valid .docx from markdown input."""

    SAMPLE_MD = """\
# Session Summary -- 2026-05-11 (Test)
**Editor**: Claude

## Francais
**Ce qui a ete fait** :
- Premiere tache accomplie
- Deuxieme tache

---

## English
**What was done**:
- First task done
- Second task
"""

    def test_output_file_created(self, tmp_path):
        src = tmp_path / "SESSION_SUMMARY.md"
        src.write_text(self.SAMPLE_MD, encoding="utf-8")
        out = tmp_path / "out.docx"
        convert(src, out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_output_contains_headings(self, tmp_path):
        from docx import Document
        src = tmp_path / "SESSION_SUMMARY.md"
        src.write_text(self.SAMPLE_MD, encoding="utf-8")
        out = tmp_path / "out.docx"
        convert(src, out)
        doc = Document(str(out))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Session Summary" in t for t in heading_texts)
        assert any("Francais" in t for t in heading_texts)
        assert any("English" in t for t in heading_texts)

    def test_output_contains_bullets(self, tmp_path):
        from docx import Document
        src = tmp_path / "SESSION_SUMMARY.md"
        src.write_text(self.SAMPLE_MD, encoding="utf-8")
        out = tmp_path / "out.docx"
        convert(src, out)
        doc = Document(str(out))
        bullet_texts = [p.text for p in doc.paragraphs if "List" in p.style.name]
        assert any("First task done" in t for t in bullet_texts)

    def test_missing_source_exits(self, tmp_path):
        src = tmp_path / "nonexistent.md"
        out = tmp_path / "out.docx"
        with pytest.raises(SystemExit):
            convert(src, out)

    def test_output_dir_created_automatically(self, tmp_path):
        src = tmp_path / "SESSION_SUMMARY.md"
        src.write_text(self.SAMPLE_MD, encoding="utf-8")
        out = tmp_path / "subdir" / "nested" / "out.docx"
        convert(src, out)
        assert out.exists()
